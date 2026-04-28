"""Convert FYDP_Report.md to a properly formatted Word .docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

MD = r"C:\Users\Muhammad Essa\Desktop\FYP\FYDP_Report.md"
OUT = r"C:\Users\Muhammad Essa\Desktop\FYP\FYDP_Report.docx"

doc = Document()

# ── Page margins ──────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1.1)
section.top_margin  = section.bottom_margin = Inches(1.0)

# ── Theme colours ─────────────────────────────────────────────
NAVY   = RGBColor(0x1a, 0x23, 0x6e)
INDIGO = RGBColor(0x43, 0x38, 0xca)
DARK   = RGBColor(0x1e, 0x1e, 0x2e)
MID    = RGBColor(0x44, 0x44, 0x66)
LIGHT  = RGBColor(0xf4, 0xf4, 0xf8)
WHITE  = RGBColor(0xff, 0xff, 0xff)
GREEN  = RGBColor(0x16, 0xa3, 0x4a)
RED    = RGBColor(0xdc, 0x26, 0x26)
CYAN   = RGBColor(0x06, 0x8f, 0xd3)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)

def add_run_with_style(para, text, bold=False, italic=False,
                       size=11, color=None, font_name="Calibri"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    return run

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Calibri"
    run.font.color.rgb = NAVY
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '1a236e')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Calibri"
    run.font.color.rgb = INDIGO
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK
    return p

def body(text):
    if not text.strip():
        return None
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    # parse inline code, bold, italic
    _add_inline(p, text)
    return p

def _add_inline(para, text):
    """Parse **bold**, *italic*, `code` inline markers and add runs."""
    # Combined pattern
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2])
            r.bold = True
            r.font.size = Pt(11)
            r.font.name = "Calibri"
            r.font.color.rgb = DARK
        elif part.startswith('*') and part.endswith('*'):
            r = para.add_run(part[1:-1])
            r.italic = True
            r.font.size = Pt(11)
            r.font.name = "Calibri"
        elif part.startswith('`') and part.endswith('`'):
            r = para.add_run(part[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xc0, 0x25, 0x25)
        else:
            r = para.add_run(part)
            r.font.size = Pt(11)
            r.font.name = "Calibri"
            r.font.color.rgb = DARK

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    _add_inline(p, text)
    return p

def code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        set_para_bg(p, 'f0f0f8')
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.right_indent = Inches(0.3)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line if line else " ")
        r.font.name = "Courier New"
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x1e, 0x40, 0x6b)

def make_table(headers, rows, header_color="1a236e"):
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Table Grid'
    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], header_color)
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = "Calibri"
        r.font.color.rgb = WHITE
    # data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        bg = 'f9f9ff' if ri % 2 == 0 else 'ffffff'
        for ci, cell_text in enumerate(row):
            set_cell_bg(cells[ci], bg)
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(cell_text))
            r.font.size = Pt(9.5)
            r.font.name = "Calibri"
            r.font.color.rgb = DARK
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(40)
r = title_p.add_run("PAYTKN")
r.bold = True
r.font.size = Pt(42)
r.font.name = "Calibri"
r.font.color.rgb = NAVY

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run("RL-Controlled Crypto Payment Token")
r2.font.size = Pt(20)
r2.font.name = "Calibri"
r2.font.color.rgb = INDIGO
r2.bold = True

doc.add_paragraph()

meta = [
    ("Final Year Design Project Report", 14, False, MID),
    ("GreyHat Finance  |  NUST NSTP Incubation Program", 13, False, MID),
    ("", 8, False, DARK),
    ("Muhammad Essa — Co-Founder, Lead Engineer", 12, True, DARK),
    ("Network: Base Sepolia Testnet  (Chain ID 84532)", 11, False, MID),
    ("Date: April 2026", 11, False, MID),
]
for txt, sz, bold, col in meta:
    mp = doc.add_paragraph()
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = mp.add_run(txt)
    mr.font.size = Pt(sz)
    mr.font.name = "Calibri"
    mr.font.color.rgb = col
    mr.bold = bold

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# PARSE MARKDOWN
# ════════════════════════════════════════════════════════════
with open(MD, encoding='utf-8', errors='replace') as f:
    lines = f.readlines()

i = 0
in_code = False
code_buf = []
in_table = False
table_headers = []
table_rows = []

def flush_table():
    global in_table, table_headers, table_rows
    if table_headers:
        make_table(table_headers, table_rows)
    in_table = False
    table_headers = []
    table_rows = []

while i < len(lines):
    line = lines[i].rstrip('\n')

    # ── Code block ───────────────────────────────────────────
    if line.strip().startswith('```'):
        if in_table:
            flush_table()
        if not in_code:
            in_code = True
            code_buf = []
        else:
            in_code = False
            code_block(code_buf)
            doc.add_paragraph()
        i += 1
        continue

    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # ── Table ────────────────────────────────────────────────
    if line.startswith('|'):
        cells = [c.strip() for c in line.strip('|').split('|')]
        if all(re.match(r'^[-:]+$', c.replace(' ', '')) for c in cells if c):
            i += 1  # skip separator row
            continue
        if not in_table:
            in_table = True
            table_headers = cells
            table_rows = []
        else:
            table_rows.append(cells)
        i += 1
        continue
    elif in_table:
        flush_table()

    # ── Images  ──────────────────────────────────────────────
    img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', line.strip())
    if img_match:
        caption = img_match.group(1)
        img_path = img_match.group(2)
        if not os.path.isabs(img_path):
            img_path = os.path.join(os.path.dirname(MD), img_path)
        if os.path.exists(img_path):
            try:
                doc.add_picture(img_path, width=Inches(5.8))
                cap_p = doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_r = cap_p.add_run(caption)
                cap_r.italic = True
                cap_r.font.size = Pt(9)
                cap_r.font.color.rgb = MID
                doc.add_paragraph()
            except Exception as e:
                body(f"[Figure: {caption}]")
        else:
            body(f"[Figure not found: {img_path}]")
        i += 1
        continue

    # ── Headings ─────────────────────────────────────────────
    if line.startswith('# ') and not line.startswith('## '):
        if line.startswith('# PAYTKN'):
            i += 1
            continue  # skip duplicate title
        heading1(line[2:].strip())
    elif line.startswith('## '):
        heading2(line[3:].strip())
    elif line.startswith('### '):
        heading3(line[4:].strip())
    elif line.startswith('#### '):
        heading3(line[5:].strip())

    # ── Bullet points ─────────────────────────────────────────
    elif line.startswith('- ') or line.startswith('* '):
        bullet(line[2:].strip())
    elif re.match(r'^\d+\. ', line):
        bullet(re.sub(r'^\d+\. ', '', line).strip())

    # ── Horizontal rule ───────────────────────────────────────
    elif line.strip() == '---':
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:color'), 'ccccdd')
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── Page break markers ─────────────────────────────────────
    elif line.strip().startswith('---') and len(line.strip()) == 3:
        pass

    # ── Regular body text ─────────────────────────────────────
    elif line.strip():
        # Skip TOC lines (they're duplicated in the doc as headings)
        if re.match(r'^\d+\. \[', line):
            pass
        else:
            body(line.strip())
    else:
        # blank line — small gap
        if i > 0 and lines[i-1].strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)

    i += 1

if in_table:
    flush_table()

# ════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════
doc.save(OUT)
print(f"Saved: {OUT}")
